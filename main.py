import os
import io
import re
import json
import shutil
import spacy
from docx import Document
from tempfile import NamedTemporaryFile
from fastapi import FastAPI, Request, File, UploadFile, Depends, HTTPException
from fastapi.responses import HTMLResponse
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
# from fastapi.templating import Jinja2Templates
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker, Session
from pdfminer.high_level import extract_text
import google.generativeai as genai
import schemas
import service
import models
import requests
from fastapi import Query

import jwt
from jwt.exceptions import InvalidTokenError

# --- Database URL ---
SQLALCHEMY_DATABASE_URL = (
    "postgresql://ryan_ai_aviation_user:riBVsmNCckogNtl9GAGKMe3dNXuKrizD"
    "@dpg-d2522j49c44c73b4f3e0-a.oregon-postgres.render.com:5432/ryan_ai_aviation"
    "?sslmode=require"
)

# --- SQLAlchemy setup ---
engine = create_engine(SQLALCHEMY_DATABASE_URL, pool_size=5, pool_timeout=20)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

app = FastAPI()
app.mount("/resumes", StaticFiles(directory="resumes"), name="resumes")
# templates = Jinja2Templates(directory="templates")

UPLOAD_FOLDER = "resumes"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

nlp = spacy.load("en_core_web_sm")
genai.configure(api_key="AIzaSyB1TMfzSPPseirr6q8KTdsRcR_ZYJtUuF4")


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# --- Helper Functions ---
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""


def extract_with_gemini(resume_text):
    prompt = f"""
    Extract the following from the resume:
    - Full Name
    - Email
    - Phone Number
    - Qualification
    - Years of Experience (as a number)
    - Career Objective
    - List of Technical Skills (comma-separated)
    - Location

    Return JSON in this format:
    {{
        "name": "...",
        "email": "...",
        "phone": "...",
        "qualification": "...",
        "experience": 3,
        "objective": "...",
        "skills": "...",
        "location": "..."
    }}

    Resume:
    \"\"\"{resume_text}\"\"\""""
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    raw_text = response.text.strip()
    cleaned = re.sub(r"^```[a-zA-Z]*\n?", "", raw_text)
    cleaned = re.sub(r"\n```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except Exception as e:
        print("Error parsing JSON:", e)
        return {}


def extract_job_keywords(description):
    prompt = f"""
    From the job description, extract:
    - Skills: list, comma separated
    - Experience: minimum years or range (e.g., "3-5", "2")

    Return JSON:
    {{
        "skills": ["skill1", "skill2", ...],
        "experience": "3-5"
    }}

    Job Description:
    \"\"\"{description}\"\"\""""
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    cleaned = re.sub(r"^```[a-zA-Z]*\n?", "", response.text.strip())
    cleaned = re.sub(r"\n```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except:
        return {"skills": [], "experience": ""}


def build_user_keywords(gpt_data):
    return {
        "skills": [s.strip().lower() for s in gpt_data.get("skills", "").split(",") if s.strip()],
        "experience": int(gpt_data.get("experience", 0))
        if str(gpt_data.get("experience", "0")).isdigit()
        else 0,
    }


def build_jobs_keywords(db: Session):
    jobs_data = db.execute(
        text('SELECT * FROM "Job" WHERE status = \'ACTIVE\'')
    ).mappings().all()

    jobs_with_keywords = []
    for job in jobs_data:
        job_keywords = extract_job_keywords(job["description"])
        jobs_with_keywords.append({
            **job,
            "skills_required": [s.lower() for s in job_keywords["skills"]],
            "experience_required": job_keywords["experience"]
        })
    return jobs_with_keywords


def compare_keywords(user_keywords, job_keywords):
    matches = []
    for job in job_keywords:
        skill_match = any(
            user_skill in job_skill or job_skill in user_skill
            for user_skill in user_keywords["skills"]
            for job_skill in job["skills_required"]
        )
        exp_match = False
        try:
            job_exp = job["experience_required"]
            if "-" in job_exp:
                low, high = job_exp.split("-")
                exp_match = int(low.strip()) <= user_keywords["experience"] <= int(
                    high.strip())
            elif job_exp.strip().isdigit():
                exp_match = user_keywords["experience"] >= int(job_exp.strip())
            else:
                exp_match = True
        except:
            exp_match = True

        if skill_match and exp_match:
            matches.append(job)
    return matches


# ------------------------------vice-versa function------------------------------

def build_candidates_keywords(db: Session):
    candidates = db.execute(
        text('SELECT "userId", "resumeUrl" FROM "CandidateProfile"')
    ).mappings().all()

    candidates_with_keywords = []
    for candidate in candidates:
        resume_url = candidate["resumeUrl"]
        try:
            # Download resume
            response = requests.get(resume_url, timeout=20)
            if response.status_code != 200:
                continue

            ext = os.path.splitext(resume_url)[-1].lower()
            resume_text = ""

            # Extract text
            if ext == ".pdf":
                resume_text = extract_text(io.BytesIO(response.content))
            elif ext == ".docx":
                doc = Document(io.BytesIO(response.content))
                resume_text = "\n".join([p.text for p in doc.paragraphs])
            else:
                continue

            # Extract structured data
            gpt_data = extract_with_gemini(resume_text)
            candidate_keywords = build_user_keywords(gpt_data)

            candidates_with_keywords.append({
                "userId": candidate["userId"],
                "resumeUrl": resume_url,
                "keywords": candidate_keywords,
                "raw_data": gpt_data
            })

        except Exception as e:
            print(f"Error processing candidate {candidate['userId']}: {e}")
            continue

    return candidates_with_keywords


def compare_job_to_candidates(job_keywords, candidates_keywords):
    matches = []

    job_skills = set(job_keywords["skills_required"])
    job_exp = job_keywords.get("experience_required", "")

    for candidate in candidates_keywords:
        user_skills = set(candidate["keywords"]["skills"])
        user_exp = candidate["keywords"].get("experience", 0)

        # --- Skill match ---
        skill_overlap = job_skills & user_skills
        skill_score = len(skill_overlap)

        # --- Experience match ---
        exp_match = False
        exp_score = 0
        try:
            if "-" in job_exp:
                low, high = job_exp.split("-")
                low, high = int(low.strip()), int(high.strip())
                exp_match = low <= user_exp <= high
                exp_score = 1 if exp_match else 0
            elif job_exp.strip().isdigit():
                required = int(job_exp.strip())
                exp_match = user_exp >= required
                exp_score = 1 if exp_match else 0
            else:
                exp_score = 1  # if not specified, ignore exp
        except:
            exp_score = 0

        # --- Final score ---
        total_score = (skill_score * 2) + exp_score

        if total_score > 0:  # at least some match
            matches.append({
                **candidate,
                "score": total_score,
                "matched_skills": list(skill_overlap),
            })

    # Sort by score (best first)
    matches.sort(key=lambda x: x["score"], reverse=True)
    return matches


def get_current_employer(request: Request):
    """Extract employer info from JWT"""
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")

    token = auth_header.split(" ")[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return payload  # will contain employer_id / role
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


# --- Routes ---
@app.get("/")
def read_root():
    return {"message": "Server is running successfully!"}


@app.get("/jobs/api", response_model=list[schemas.job])
def get_jobs_api(db: Session = Depends(get_db)):
    return service.get_job(db)


@app.get("/jobs")
async def index(request: Request, db: Session = Depends(get_db)):
    jobs = service.get_job(db)
    # return templates.TemplateResponse("index.html", {"request": request, "all_jobs": jobs})
    return jobs


SECRET_KEY = "TK2fmxH/OhyoC5M1nBc1shbw5xvjtrTOq1lSBJ1svR2BXgJkuNrSYnSPPVw="
ALGORITHM = "HS256"


@app.get("/recommend_jobs")
async def recommend_jobs_from_token(request: Request, db: Session = Depends(get_db)):
    # 1. Get JWT token from Authorization header
    auth_header = request.headers.get("Authorization")
    # print("Auth Header:", auth_header)
    if not auth_header or not auth_header.startswith("Bearer "):
        return JSONResponse({"error": "Missing or invalid token"}, status_code=401)
    token = auth_header.split(" ")[1]

    # 2. Decode the token
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        print("found user_id:", user_id)
        if not user_id:
            return JSONResponse({"error": "User ID not found in token"}, status_code=400)
    except InvalidTokenError:
        return JSONResponse({"error": "Invalid token"}, status_code=401)

    # 3. Get resume URL from DB
    result = db.execute(
        text('SELECT "resumeUrl" FROM "CandidateProfile" WHERE "userId" = :id'),
        {"id": user_id}
    ).mappings().first()

    if not result or not result["resumeUrl"]:
        return JSONResponse({"error": "Resume URL not found for user"}, status_code=404)

    resume_url = result["resumeUrl"]

    # 4. Download the file
    try:
        response = requests.get(resume_url, timeout=20)
        if response.status_code != 200:
            return JSONResponse({"error": "Unable to download resume"}, status_code=400)

        # Save temp file
        temp_path = os.path.join(UPLOAD_FOLDER, f"user_{user_id}_resume")
        ext = os.path.splitext(resume_url)[-1].lower()
        with open(temp_path + ext, "wb") as f:
            f.write(response.content)

        # Extract text
        if ext == ".pdf":
            resume_text = extract_text(temp_path + ext)
        elif ext == ".docx":
            resume_text = extract_text_from_docx(temp_path + ext)
        else:
            return JSONResponse({"error": "Unsupported file format"}, status_code=400)

    except Exception as e:
        return JSONResponse({"error": f"Error processing resume: {str(e)}"}, status_code=500)

    # 5. Extract structured data using Gemini
    gpt_data = extract_with_gemini(resume_text)

    # 6. Build user keywords
    user_keywords = build_user_keywords(gpt_data)

    # 7. Build job keywords from DB
    jobs_with_keywords = build_jobs_keywords(db)

    # 8. Compare and find matches
    recommended_jobs = compare_keywords(user_keywords, jobs_with_keywords)

    return {"recommended_jobs": recommended_jobs}


# -------------------------------------vice-versa api-------------------------------------

@app.get("/recommend_candidates/{job_id}")
async def recommend_candidates_for_job(
    job_id: str,   # cuid/uuid style for Job.id
    request: Request,
    db: Session = Depends(get_db)
):
    # --- Get JWT token ---
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return JSONResponse({"error": "Missing or invalid token"}, status_code=401)

    token = auth_header.split(" ")[1]
    try:
        decoded_token = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
    except jwt.ExpiredSignatureError:
        return JSONResponse({"error": "Token expired"}, status_code=401)
    except jwt.InvalidTokenError:
        return JSONResponse({"error": "Invalid token"}, status_code=401)

    # --- Extract user_id from token ---
    user_id = decoded_token.get("sub")
    if not user_id:
        return JSONResponse({"error": "User ID not found in token"}, status_code=400)

    # ✅ Fetch employer_id from DB using user_id (no type mismatch now)
    employer = db.execute(
        text('SELECT id FROM "EmployerProfile" WHERE "userId" = :uid'),
        {"uid": user_id}
    ).mappings().first()

    if not employer:
        return JSONResponse({"error": "Employer not found"}, status_code=404)

    employer_id = employer["id"]

    # ✅ 2. Verify job belongs to this employer directly in query
    job = db.execute(
        text("""
            SELECT * 
            FROM "Job" 
            WHERE id = :jid 
              AND "employerId" = :eid 
              AND status = 'ACTIVE'
        """),
        {"jid": job_id, "eid": employer_id}
    ).mappings().first()

    if not job:
        return JSONResponse({"error": "Job not found or not owned by this employer"}, status_code=404)

    # ✅ 3. Extract job keywords
    job_keywords = extract_job_keywords(job["description"])
    job_with_keywords = {
        **job,
        "skills_required": [s.lower() for s in job_keywords["skills"]],
        "experience_required": job_keywords["experience"]
    }

    # ✅ 4. Get all candidates
    candidates_with_keywords = build_candidates_keywords(db)

    # ✅ 5. Compare job → candidates
    recommended_candidates = compare_job_to_candidates(
        job_with_keywords, candidates_with_keywords
    )

    # ✅ 6. Build output
    result = []
    for candidate in recommended_candidates:
        profile = db.execute(
            text("""
                SELECT cp.id,
                       cp."jobCategory",
                       cp."currentLocation",
                       cp."totalExperience",
                       cp."nationality",
                       cp."resumeUrl",
                       u."fullName",
                       u."image"
                FROM "CandidateProfile" cp
                JOIN "User" u ON cp."userId" = u.id
                WHERE cp."userId" = :uid
            """),
            {"uid": candidate["userId"]}
        ).mappings().first()

        if not profile:
            continue

        bookmark = db.execute(
            text(
                'SELECT 1 FROM "CandidateBookmark" WHERE "employerId" = :eid AND "candidateId" = :cid'
            ),
            {"eid": employer_id, "cid": profile["id"]}
        ).first()

        result.append({
            "id": profile["id"],
            "fullName": profile["fullName"],
            "image": profile["image"],
            "jobCategory": profile["jobCategory"],
            "currentLocation": profile["currentLocation"],
            "totalExperience": profile["totalExperience"],
            "nationality": profile["nationality"],
            "resumeUrl": profile["resumeUrl"],
            "isBookmarked": bool(bookmark)
        })

    return {"recommended_candidates": result}
